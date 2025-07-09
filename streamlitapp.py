import streamlit as st
from azure.cognitiveservices.vision.computervision import ComputerVisionClient
from msrest.authentication import CognitiveServicesCredentials
from PIL import Image
import io
import openai
import os
import time
from docx import Document

st.set_page_config(page_title="Gerador de Relatório STRIDE", layout="centered")
st.title("📊 Gerador de Relatório STRIDE")

st.markdown("### 🔐 Credenciais da Azure Computer Vision")
endpoint = st.text_input("Azure Endpoint")
key = st.text_input("Azure Key", type="password")

st.markdown("### 🔐 Credenciais da OpenAI")
openai_api_key = st.text_input("OpenAI API Key", type="password")

uploaded_file = st.file_uploader("📁 Upload do Diagrama de Arquitetura", type=["png", "jpg", "jpeg"])

# Only process if not already done and inputs are valid
if uploaded_file and endpoint and key and openai_api_key:

    if "report_text" not in st.session_state:
        st.success("Processando a imagem...")

        computervision_client = ComputerVisionClient(
            endpoint, CognitiveServicesCredentials(key)
        )

        # Read image
        image_bytes = uploaded_file.read()
        image_stream = io.BytesIO(image_bytes)

        ocr_result = computervision_client.read_in_stream(image_stream, raw=True)
        operation_location = ocr_result.headers["Operation-Location"]
        operation_id = operation_location.split("/")[-1]

        while True:
            result = computervision_client.get_read_result(operation_id)
            if result.status not in ['notStarted', 'running']:
                break
            time.sleep(1)

        extracted_text = ""
        if result.status == 'succeeded':
            for page in result.analyze_result.read_results:
                for line in page.lines:
                    extracted_text += line.text + "\n"
        else:
            st.error("❌ OCR failed.")
            st.stop()

        st.session_state.extracted_text = extracted_text

        # Prompt Engineering
        prompt = f"""
You are a system architect. Take all the necessary time to create a quality response and with as much detail as possible. Validate if the answer is not contraditory in any capacity. Bring references to each of the topics. Do not be long-winded. 
Do not generate any observation or follow-up commentary after the STRIDE report. 
Based on the following extracted architecture diagram text, write a STRIDE report in brazillian portuguese, adding the risks, mitigation and references for the following sections:

1. **Spoofing**
2. **Tampering**
3. **Repudiation**
4. **Information Disclosure**
5. **Denial of Service**
6. **Elevation of Privilege**

After the STRIDE Report, generate a table in markdown, with all the stride sections in one column, risks in the second column and mitigation in the third column. Be very concise on the second and third columns.

Extracted Text:
{extracted_text}
"""

        client = openai.OpenAI(api_key=openai_api_key)
        response = client.chat.completions.create(
            model="gpt-4.1",
            messages=[
                {"role": "system", "content": "You are a system architecture analyst."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4,
        )

        report_text = response.choices[0].message.content
        st.session_state.report_text = report_text

        # Generate Word document
        doc = Document()
        doc.add_heading('Relatório Stride', 0)
        for line in report_text.split('\n'):
            if line.startswith("**") and line.endswith("**"):
                doc.add_heading(line.replace("**", ""), level=1)
            else:
                doc.add_paragraph(line)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.session_state.doc_buffer = buffer

# If already processed, show results
if "report_text" in st.session_state:
    st.markdown("### 📝 Texto Extraído")
    st.text_area("OCR Output", st.session_state.extracted_text, height=200)

    st.markdown("### 🧾 Relatório Stride")
    st.markdown(st.session_state.report_text)

    st.download_button(
        label="📄 Download Relatório Stride (.docx)",
        data=st.session_state.doc_buffer,
        file_name="Relatorio_Stride.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

elif uploaded_file:
    st.info("O processamento inicializará assim que todas as credenciais forem submetidas")
else:
    st.warning("Por favor, coloque as credenciais e carregue a imagem do diagrama")
